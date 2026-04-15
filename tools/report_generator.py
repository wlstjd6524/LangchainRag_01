"""
report_generator.py
ESG 에이전트 보고서 생성 툴

LLM이 분석한 결과를 Word(.docx) 형식의 보고서로 저장한다.
"""

import os
from datetime import datetime
from langchain_core.tools import tool
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPORTS_DIR = "./reports"


def _ensure_reports_dir():
    os.makedirs(REPORTS_DIR, exist_ok=True)


def _add_heading(doc: Document, text: str, level: int = 1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)  # ESG 녹색


def _build_docx(title: str, content: str) -> Document:
    doc = Document()

    # 제목
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

    # 생성 일시
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"생성 일시: {datetime.now().strftime('%Y년 %m월 %d일 %H:%M')}")
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)

    doc.add_paragraph()  # 여백

    # 본문 파싱: "## 제목" → heading, 나머지 → 본문
    for line in content.splitlines():
        stripped = line.strip()
        if stripped.startswith("## "):
            _add_heading(doc, stripped[3:], level=2)
        elif stripped.startswith("# "):
            _add_heading(doc, stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            para = doc.add_paragraph(stripped[2:], style="List Bullet")
            para.runs[0].font.size = Pt(11)
        elif stripped == "":
            doc.add_paragraph()
        else:
            para = doc.add_paragraph(stripped)
            para.runs[0].font.size = Pt(11)

    return doc


@tool
def generate_report(title: str, content: str) -> str:
    """
    사용자가 분석 결과나 답변을 보고서로 저장하고 싶을 때 호출한다.
    사용자가 이번 메시지에서 "저장", "파일", "보고서", "워드", "word", "보고서로 저장해줘", "파일로 만들어줘", "Word로 저장해줘"등과 같이 하나 이상의 단어로 명시적으로 파일 저장을 요청한 경우에만 호출한다.
    분석 또는 점검 결과, 검색 결과가 나왔다고 해서 이 툴을 자동으로 호출하지 않는다.
    
    다음 상황에서는 절대 호출하지 않는다:
    - 분석, 계산, 검색 결과가 출력된 직후
    - 사용자가 "분석해줘", "확인해줘", "분석을 해줘"라고 요청한 경우 (보고서 저장이 아니라 분석 자체를 요청한 경우)
    - 파일 저장을 명시적으로 요청하지 않은 모든 상황

    이 툴은 반드시 1회만 호출하며, 모든 내용을 content에 담아 한 번에 저장한다.

    Args:
        title: 보고서 제목 (예: "ESG 안전지표 분석 보고서")
        content: 보고서 본문 내용 전체. 마크다운 형식 지원 (## 제목, - 항목 등)

    Returns:
        저장된 파일 경로
    """
    print("##### REPORT GENERATOR TOOL #####")

    try:
        _ensure_reports_dir()

        safe_title = "".join(c for c in title if c.isalnum() or c in (" ", "_")).strip()[:30]

        # 같은 제목의 파일이 이미 있으면 덮어씀 (중복 방지)
        existing = [f for f in os.listdir(REPORTS_DIR) if f.endswith(f"_{safe_title}.docx")]
        if existing:
            filepath = os.path.join(REPORTS_DIR, existing[0])
        else:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filepath = os.path.join(REPORTS_DIR, f"{timestamp}_{safe_title}.docx")

        doc = _build_docx(title, content)
        doc.save(filepath)

        return f"보고서가 저장되었습니다.\n파일 경로: {os.path.abspath(filepath)}"

    except Exception as e:
        return f"보고서 저장 오류: {str(e)}"
